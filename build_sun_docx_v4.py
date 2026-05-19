"""
生成孙晓燕简历的 Word 版本（含证件照）
"""
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ----- 主题色 -----
NAVY = RGBColor(0x1E, 0x3A, 0x5F)         # 深蓝主色
SUBTITLE_BG = RGBColor(0x4A, 0x6B, 0x8A)  # 二级标题底色
TITLE_BG = NAVY                            # 一级标题底色
LIGHT_BG = RGBColor(0xF4, 0xF7, 0xFB)     # 浅蓝底
TEXT_DARK = RGBColor(0x2C, 0x3E, 0x50)
TEXT_GREY = RGBColor(0x55, 0x55, 0x55)
TEXT_GREY_LIGHT = RGBColor(0x77, 0x77, 0x77)
TEXT_DARK_GREY = RGBColor(0x44, 0x44, 0x44)
TEXT_LIGHT_GREY = RGBColor(0x66, 0x66, 0x66)


def set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_paragraph_bg(paragraph, hex_color: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def set_paragraph_left_border(paragraph, hex_color='1E3A5F', size='18'):
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), hex_color)
    pBdr.append(left)
    p_pr.append(pBdr)


def set_paragraph_bottom_border(paragraph, hex_color='1E3A5F', size='18'):
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    p_pr.append(pBdr)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'nil')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_run(paragraph, text, bold=False, size=10, color=None, font='等线'):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    # 中文字体设置
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    if color is not None:
        run.font.color.rgb = color
    return run


def set_paragraph_spacing(p, before=0, after=2, line=1.4):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def section_title(doc, text):
    """添加深蓝底白字一级标题"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=3, line=1.3)
    set_paragraph_bg(p, '1E3A5F')
    pf = p.paragraph_format
    pf.left_indent = Cm(0)
    add_run(p, ' ' + text + ' ', bold=True, size=12, color=RGBColor(0xFF, 0xFF, 0xFF))
    return p


def sub_title(doc, text):
    """二级标题：浅蓝底白字"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=1, line=1.3)
    set_paragraph_bg(p, '4A6B8A')
    add_run(p, ' ' + text + ' ', bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
    return p


def bullet(doc, label, content, sub_bold_keywords=None):
    """带 ▸ 符号的项目，label 加粗"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=1, line=1.4)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    add_run(p, '▸ ', bold=True, size=10, color=NAVY)
    add_run(p, label, bold=True, size=10, color=NAVY)
    add_run(p, ' ｜ ', size=10, color=TEXT_GREY)
    # 处理内容里的加粗关键词
    if sub_bold_keywords:
        rest = content
        for kw in sub_bold_keywords:
            idx = rest.find(kw)
            if idx >= 0:
                add_run(p, rest[:idx], size=10, color=TEXT_DARK)
                add_run(p, kw, bold=True, size=10, color=NAVY)
                rest = rest[idx+len(kw):]
        add_run(p, rest, size=10, color=TEXT_DARK)
    else:
        add_run(p, content, size=10, color=TEXT_DARK)
    return p


def simple_bullet(doc, content_runs):
    """通用 ▸ 项目，content_runs 是 (text, bold) 列表"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=1, line=1.4)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    add_run(p, '▸ ', bold=True, size=10, color=NAVY)
    for text, bold in content_runs:
        add_run(p, text, bold=bold, size=10, color=NAVY if bold else TEXT_DARK)
    return p


def main():
    doc = Document()

    # 页面边距
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    # A4
    section.page_width = Mm(210)
    section.page_height = Mm(297)

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '等线'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    # ========== 头部：左姓名+联系方式，右照片 ==========
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    remove_table_borders(header_table)
    header_table.columns[0].width = Cm(13.5)
    header_table.columns[1].width = Cm(3.2)
    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]
    left_cell.width = Cm(13.5)
    right_cell.width = Cm(3.2)

    # 左侧：姓名
    name_p = left_cell.paragraphs[0]
    set_paragraph_spacing(name_p, before=0, after=4, line=1.2)
    name_run = add_run(name_p, '孙  晓  燕', bold=True, size=26, color=NAVY)

    # 联系方式 -- 用两行
    contact_p1 = left_cell.add_paragraph()
    set_paragraph_spacing(contact_p1, before=2, after=2, line=1.4)
    add_run(contact_p1, '出生: ', bold=True, size=9.5, color=NAVY)
    add_run(contact_p1, '1989.11      ', size=9.5, color=TEXT_GREY)
    add_run(contact_p1, '学历: ', bold=True, size=9.5, color=NAVY)
    add_run(contact_p1, '本科      ', size=9.5, color=TEXT_GREY)
    add_run(contact_p1, '现居: ', bold=True, size=9.5, color=NAVY)
    add_run(contact_p1, '青岛市黄岛区      ', size=9.5, color=TEXT_GREY)
    add_run(contact_p1, '籍贯: ', bold=True, size=9.5, color=NAVY)
    add_run(contact_p1, '山东青岛', size=9.5, color=TEXT_GREY)

    contact_p2 = left_cell.add_paragraph()
    set_paragraph_spacing(contact_p2, before=2, after=2, line=1.4)
    add_run(contact_p2, '电话/微信: ', bold=True, size=9.5, color=NAVY)
    add_run(contact_p2, '13370819121      ', size=9.5, color=TEXT_GREY)
    add_run(contact_p2, '邮箱: ', bold=True, size=9.5, color=NAVY)
    add_run(contact_p2, 'cindy_sun0@163.com', size=9.5, color=TEXT_GREY)

    # 右侧：照片
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    photo_p = right_cell.paragraphs[0]
    photo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(photo_p, before=0, after=0, line=1.0)
    photo_run = photo_p.add_run()
    photo_run.add_picture('/Users/tiny/Desktop/resume/avatar_xiaoyan.jpg', width=Cm(2.8), height=Cm(3.6))

    # 头部底部分隔线
    sep = doc.add_paragraph()
    set_paragraph_spacing(sep, before=0, after=0, line=0.5)
    set_paragraph_bottom_border(sep, '1E3A5F', '24')

    # ========== 个人优势 ==========
    section_title(doc, '个 人 优 势')
    intro = doc.add_paragraph()
    set_paragraph_spacing(intro, before=2, after=4, line=1.5)
    set_paragraph_bg(intro, 'F4F7FB')
    set_paragraph_left_border(intro, '1E3A5F', '24')
    intro.paragraph_format.left_indent = Cm(0.3)
    intro_text = (
        '10 余年外资制造业物流关务经验，深耕世界 500 强德资企业及大型显示面板龙头，'
        '精通进出口物流（海/陆/空运）、关务申报与合规、AEO 体系建设、订单与供应商管理。'
        '德语专业八级（PGH），英语六级且可作为工作语言，能独立对接海外总部、海关与外籍客户。'
        '熟悉 SAP 主数据、加工贸易、保税监管、出口退税及鼓励类减免税政策，'
        '曾主导多项流程优化和成本节约项目，具备跨部门协同与海关稽查应对的实战经验。'
    )
    add_run(intro, intro_text, bold=False, size=10, color=TEXT_DARK)

    # ========== 工作经历 ==========
    section_title(doc, '工 作 经 历')

    # ---- 京东方 ----
    jh = doc.add_paragraph()
    set_paragraph_spacing(jh, before=4, after=1, line=1.3)
    set_paragraph_bottom_border(jh, '1E3A5F', '14')
    add_run(jh, '青岛京东方光电科技有限公司', bold=True, size=11, color=NAVY)
    add_run(jh, '   物流关务担当', bold=True, size=10.5, color=NAVY)
    add_run(jh, '\t', size=10)
    jh.paragraph_format.tab_stops.add_tab_stop(Cm(16))
    add_run(jh, '2023.09 – 至今 ｜ 2 年+', bold=True, size=9.5, color=NAVY)

    tag = doc.add_paragraph()
    set_paragraph_spacing(tag, before=1, after=2, line=1.3)
    set_paragraph_bg(tag, 'F4F7FB')
    add_run(tag, '  行业：显示面板行业龙头（A 股上市） ｜ 团队：物流关务部 ｜ 汇报对象：部门经理  ',
            size=9, color=TEXT_LIGHT_GREY)

    summary = doc.add_paragraph()
    set_paragraph_spacing(summary, before=2, after=2, line=1.5)
    add_run(summary, '岗位职责：', bold=True, size=10, color=NAVY)
    add_run(summary,
            '负责京东方青岛工厂设备/备件/原材料/消耗品的进出口物流与通关全流程，'
            '统筹关务合规、AEO 持续符合管理、海关对接及跨部门审核支持。',
            size=10, color=TEXT_DARK_GREY)

    sub_title(doc, '物 流 与 通 关 操 作')
    bullet(doc, '进口物流统筹', '安排海/陆/空运及国外直购、保税园进口模式，持续推动流程优化与成本节约。')
    bullet(doc, '设备搬运管理', '跨厂区设备搬迁、公司内部工具吊装；监督供应商符合 EHS 及合规要求。')
    bullet(doc, '运输保险', '主导集团内部分区域货物的投保、付款与理赔，保障货物风险可控。')

    sub_title(doc, '关 务 合 规 管 理')
    bullet(doc, '申报规范', '审核商品归类、维护 SAP 主数据，宣导申报变更并处置关务异常。')
    bullet(doc, '报关行管理', '资质审核 + 报关授权 + KPI 量化考核，提升服务商质量。')
    bullet(doc, '数据分析', '维护进出口报表，分析进出口额与税款变动，为业务决策提供数据支撑。')
    bullet(doc, '政策应用', '跟踪法规更新，运用贸易协定、原产地规则、税收优惠降本并防范风险。')
    bullet(doc, '海关对接', '应对稽核查、贸易调查等，主动参与关企沟通活动。')
    bullet(doc, 'AEO 持续符合', '组织合规内审、开展全员 AEO 培训、年度 AEO 内审初审与资料归档、按时报送企业年报。')
    bullet(doc, '跨部门协同', '配合供应链审核、质量体系审核、内控与财务审计等多类审核。')

    # ---- 德枫丹 ----
    jh2 = doc.add_paragraph()
    set_paragraph_spacing(jh2, before=6, after=1, line=1.3)
    set_paragraph_bottom_border(jh2, '1E3A5F', '14')
    add_run(jh2, '德枫丹（青岛）机械有限公司', bold=True, size=11, color=NAVY)
    add_run(jh2, '   物流专员', bold=True, size=10.5, color=NAVY)
    add_run(jh2, '\t', size=10)
    jh2.paragraph_format.tab_stops.add_tab_stop(Cm(16))
    add_run(jh2, '2014.09 – 2023.09 ｜ 9 年', bold=True, size=9.5, color=NAVY)

    tag2 = doc.add_paragraph()
    set_paragraph_spacing(tag2, before=1, after=2, line=1.3)
    set_paragraph_bg(tag2, 'F4F7FB')
    add_run(tag2, '  行业：世界 500 强 · 德国蒂森克虏伯 ThyssenKrupp 集团全资子公司 ｜ 业务：高端机电设备制造  ',
            size=9, color=TEXT_LIGHT_GREY)

    summary2 = doc.add_paragraph()
    set_paragraph_spacing(summary2, before=2, after=2, line=1.5)
    add_run(summary2, '岗位职责：', bold=True, size=10, color=NAVY)
    add_run(summary2,
            '独立负责公司进出口物流与关务全链路——从供应商管理、订单运输、通关申报到政策应用、'
            'AEO 体系与单证存档；服务期内多次实现关税减免与物流降本。',
            size=10, color=TEXT_DARK_GREY)

    sub_title(doc, '物 流 与 供 应 商')
    bullet(doc, '供应商管理', '筛选报关行与第三方物流商，议定海/空/铁运价并定期评估，降低报关差错率、提升送达及时率。')
    bullet(doc, '订单与运输', '跟踪生产进度，及时订舱并选择最优运输方案；维护发货跟踪表，处置运输异常。')
    bullet(doc, '包装与成本', '参与包装方案设计、编制月度物流与保险费用报告、关注海运行情，持续优化降本。')

    sub_title(doc, '通 关 与 证 照')
    bullet(doc, '通关业务', '熟悉新旧机电产品、设备备件、原材料、化学品及危险品的进出口流程与商品归类。')
    bullet(doc, '证照办理', '进出口许可证、免 3C、装运前检验、化学品鉴定、原产地证等全套证书申请。')
    bullet(doc, '加工贸易', '熟悉保税物料管理；配合财务完成出口退税与增值税发票开具。')

    sub_title(doc, '政 策 应 用 与 合 规 体 系')
    bullet(doc, '政策红利', '★ 成功为公司申请鼓励类产业项目确认书，实现设备进口关税减免，显著节省成本。')
    bullet(doc, '后续监管', '减免税设备年报报送、免 3C 货物核销、台账记录等特殊监管货物全周期跟踪。')
    bullet(doc, '体系建设', '参与 AEO 认证；起草并完善进出口文件，优化流程并规范单证存档。')

    # ---- 信高夫 ----
    jh3 = doc.add_paragraph()
    set_paragraph_spacing(jh3, before=6, after=1, line=1.3)
    set_paragraph_bottom_border(jh3, '1E3A5F', '14')
    add_run(jh3, '青岛信高夫进出口有限公司', bold=True, size=11, color=NAVY)
    add_run(jh3, '   外贸专员', bold=True, size=10.5, color=NAVY)
    add_run(jh3, '\t', size=10)
    jh3.paragraph_format.tab_stops.add_tab_stop(Cm(16))
    add_run(jh3, '2013.07 – 2014.09 ｜ 1 年+', bold=True, size=9.5, color=NAVY)

    tag3 = doc.add_paragraph()
    set_paragraph_spacing(tag3, before=1, after=2, line=1.3)
    set_paragraph_bg(tag3, 'F4F7FB')
    add_run(tag3, '  行业：进出口贸易 ｜ 主要市场：德国 / 欧洲  ', size=9, color=TEXT_LIGHT_GREY)

    bullet(doc, '订单管理', '分析客户询价、与工厂确认交期并报价、签订合同、传递订单要求至工厂。')
    bullet(doc, '物流协同', '跟踪生产进度，与货代、保险公司协调订舱发货；跟踪并确认客户付款。')
    bullet(doc, '德语支持', '陪同接待来访德国客户，承担德语类文件翻译，为商务谈判提供语言支持。')

    # ========== 教育背景 ==========
    section_title(doc, '教 育 背 景')
    edu = doc.add_paragraph()
    set_paragraph_spacing(edu, before=2, after=1, line=1.4)
    add_run(edu, '济南大学', bold=True, size=10.5, color=NAVY)
    add_run(edu, '  |  ', size=10, color=TEXT_GREY)
    add_run(edu, '德语专业（本科）', bold=True, size=10, color=TEXT_DARK)
    edu.paragraph_format.tab_stops.add_tab_stop(Cm(16))
    add_run(edu, '\t', size=10)
    add_run(edu, '2009.09 – 2013.06', size=9.5, color=TEXT_GREY_LIGHT)

    edu2 = doc.add_paragraph()
    set_paragraph_spacing(edu2, before=1, after=2, line=1.5)
    add_run(edu2, '荣誉奖项：', bold=True, size=9.5, color=NAVY)
    add_run(edu2, '国家励志奖学金 1 次 · 校二等奖学金 2 次 · 校三等奖学金 1 次 · "优秀团员"称号 2 次',
            size=9.5, color=TEXT_GREY)

    # ========== 证书与技能 ==========
    section_title(doc, '证 书 与 技 能')

    rows = [
        ('语言能力：', '德语专业八级（PGH） · 英语六级（CET-6），可作为工作语言对接外籍客户与海外总部'),
        ('职业资格：', '二级企业人力资源管理师'),
    ]
    for label, content in rows:
        sp = doc.add_paragraph()
        set_paragraph_spacing(sp, before=1, after=1, line=1.5)
        add_run(sp, label, bold=True, size=10, color=NAVY)
        add_run(sp, content, bold=False, size=10, color=TEXT_DARK)

    out = '/Users/tiny/Desktop/resume/孙晓燕_物流关务简历_含证件照_v2.docx'
    doc.save(out)
    print(f'saved -> {out}')


if __name__ == '__main__':
    main()
